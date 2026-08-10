import mysql.connector
from docx import Document

def clear():
    print("\033c")
    
def pause():
    input("... ¡Oprima cualquier tecla para continuar! ...")
    
def finish():
    clear()
    input("....::::¡SESION FINALIZADA CON EXITO!::::....")
    
def invalid():
    input("\n\t ..::¡Opcion inválida!::..\n\n Oprima cualquier tecla para continuar....")

def success():
    input("\n\t...¡Accion Realizada con Exito!...")
    
def unsuccess():
    input("\n\t...¡Esta accion no pudo ser realizada!...")
    
def connection():
    try:
        conexion=mysql.connector.connect(
            host="127.0.0.1",
            user="root",
            password="",
            database="db_gym"
        )
        return conexion
    except:
        clear()
        input("... ¡Por el momento no es posible establecer conexion con la base de datos. Por favor intententelo mas tarde! ...") 
        return None   

def mainMenu():
    print("\n\t\t\t...::: M E N U  P R I N C I P A L :::... \n")
    opcion=input("\n\t 1.- Gestión de clientes \n\t 2.- Gestión de membresías \n\t 3.- Salir\n\n\t\tSeleccionar Opcion: ").strip()
    return opcion

def repClientes(customers):
    archivo = open("Reporte de clientes.txt", "w", encoding="utf-8")
    archivo.write("REPORTE DE CLIENTES")
    archivo.write("=" * 60 + "")
    for i in customers:
        archivo.write( f"ID: {i[0]}\nNombre: {i[1]}\nApellido: {i[2]}\nTeléfono: {i[3]}\nCorreo: {i[4]}" )
        archivo.close()

def repClientesDocx(customers):
    doc = Document()
    doc.add_heading('Reporte de clientes', level=1)
    for i in customers:
        doc.add_paragraph( f"ID: {i[0]}\nNombre: {i[1]}\nApellido: {i[2]}\nTeléfono: {i[3]}\nCorreo: {i[4]}" )
    doc.save("Reporte de clientes.docx")


    